"""Word document creation and editing."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from utils.helpers import ensure_directory, safe_filename


class WordController:
    """Create, append to, and read Word documents with python-docx."""

    def __init__(self, config: Any) -> None:
        self.config = config
        self.directory = ensure_directory(config.get("paths.documents_dir", "documents"))
        self.current_doc: Path | None = None

    def create(self, title: str) -> str:
        filename = safe_filename(title, "jarvis_document") + ".docx"
        path = self.directory / filename
        try:
            from docx import Document  # type: ignore

            doc = Document()
            doc.add_heading(title, level=1)
            doc.save(path)
        except Exception:
            path = path.with_suffix(".txt")
            path.write_text(title + "\n\n", encoding="utf-8")
        self.current_doc = path
        return f"Document created at {path}, sir."

    def append(self, text: str) -> str:
        if not self.current_doc:
            self.create("Jarvis Notes")
        assert self.current_doc is not None
        if self.current_doc.suffix == ".docx":
            try:
                from docx import Document  # type: ignore

                doc = Document(self.current_doc)
                doc.add_paragraph(text)
                doc.save(self.current_doc)
                return "Text added to the document, sir."
            except Exception:
                pass
        with self.current_doc.open("a", encoding="utf-8") as handle:
            handle.write(text + "\n")
        return "Text added to the document, sir."

    def read(self) -> str:
        if not self.current_doc or not self.current_doc.exists():
            return "No active document is open, sir."
        if self.current_doc.suffix == ".docx":
            try:
                from docx import Document  # type: ignore

                doc = Document(self.current_doc)
                text = " ".join(p.text for p in doc.paragraphs if p.text).strip()
                return text[:1200] if text else "The document is empty, sir."
            except Exception:
                return "I could not read the Word document, sir."
        return self.current_doc.read_text(encoding="utf-8")[:1200]

    def process(self, command: str) -> dict[str, Any]:
        lower = command.lower()
        if "document" not in lower and "word" not in lower:
            return {"success": False, "response": "I did not find a Word command, sir."}
        create_match = re.search(r"create (?:word )?document(?: called| named)? (.+)", command, re.I)
        if create_match:
            return {"success": True, "response": self.create(create_match.group(1).strip())}
        append_match = re.search(r"(?:add|append|write) (?:to )?(?:document|word)[: ]+(.+)", command, re.I | re.S)
        if append_match:
            return {"success": True, "response": self.append(append_match.group(1).strip())}
        if "read" in lower or "summarize" in lower:
            return {"success": True, "response": self.read()}
        return {"success": False, "response": "I did not find a Word command, sir."}
